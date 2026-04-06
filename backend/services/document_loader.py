"""
Document Loader Service
Extracts text from PDF, DOCX, XLSX, and TXT files.
"""
import os
import time
from pathlib import Path
from typing import List, Dict
import openpyxl
from docx import Document as DocxDocument
from pypdf import PdfReader
from PIL import Image
import google.generativeai as genai


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF. Fallback to Gemini OCR if it's a scanned document."""
    # 1. Try standard text extraction
    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    
    text = "\n\n".join(text_parts)
    
    # 2. If text is missing or sparse, it's a scanned PDF. Use Bedrock Vision OCR.
    if len(text.strip()) < 100:
        print(f"[Loader] Low text ({len(text.strip())} chars) in {Path(file_path).name}. Trying Bedrock OCR...")
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            ocr_parts = []
            for i, page in enumerate(doc):
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                try:
                    page_text = _ocr_image_bytes(img_bytes, "png", f"page {i+1}")
                    if page_text.strip():
                        ocr_parts.append(page_text)
                except Exception as e:
                    print(f"[Loader] OCR failed page {i+1}: {e}")
            doc.close()
            if ocr_parts:
                return "\n\n".join(ocr_parts)
        except ImportError:
            print("[Loader] PyMuPDF not installed. Run: pip install PyMuPDF")
        except Exception as e:
            print(f"[Loader] Bedrock OCR fallback failed: {e}")
    
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file."""
    doc = DocxDocument(file_path)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                text_parts.append(" | ".join(row_data))

    return "\n\n".join(text_parts)


def extract_text_from_xlsx(file_path: str) -> str:
    """Extract all text from an Excel file, treating each row as a paragraph."""
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"--- Sheet: {sheet_name} ---")

        headers = []
        for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
            values = [str(cell) if cell is not None else "" for cell in row]
            non_empty = [v for v in values if v.strip()]
            if not non_empty:
                continue

            if row_idx == 0:
                headers = values
                text_parts.append("Headers: " + " | ".join(non_empty))
            else:
                # Combine header+value pairs for richer context
                if headers:
                    pairs = []
                    for h, v in zip(headers, values):
                        if v.strip():
                            pairs.append(f"{h}: {v}" if h.strip() else v)
                    if pairs:
                        text_parts.append(" | ".join(pairs))
                else:
                    text_parts.append(" | ".join(non_empty))

    wb.close()
    return "\n".join(text_parts)


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def _ocr_image_bytes(image_bytes: bytes, fmt: str, label: str = "") -> str:
    """OCR image bytes using Groq Vision (Llama 3.2 11B)."""
    import base64, requests
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    media_types = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}
    media_type = media_types.get(fmt, "image/jpeg")

    resp = requests.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={"Authorization": f"Bearer {os.getenv('GROQ_API_KEY')}", "Content-Type": "application/json"},
        json={
            "model": "meta-llama/llama-4-scout-17b-16e-instruct",
            "messages": [{"role": "user", "content": [
                {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{b64}"}},
                {"type": "text", "text": "Extract ALL text from this image exactly as written. Preserve formatting."}
            ]}],
            "max_tokens": 2000,
        },
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


def extract_text_from_image(file_path: str) -> str:
    """Extract text from an image using Groq Vision (free, no AWS needed)."""
    import base64

    with open(file_path, "rb") as f:
        image_bytes = f.read()

    suffix = Path(file_path).suffix.lower().lstrip(".")
    text = _ocr_image_bytes(image_bytes, suffix, Path(file_path).name)
    return f"--- Image Content ({Path(file_path).name}) ---\n{text}"

def extract_text_from_video(file_path: str) -> str:
    """Upload video to Gemini and describe it fully."""
    genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
    video_file = genai.upload_file(path=file_path)
    
    # Wait for processing with a 5-minute timeout
    max_wait = 300
    start_time = time.time()
    while video_file.state.name == "PROCESSING":
        if time.time() - start_time > max_wait:
            raise TimeoutError(f"Gemini Video processing timed out for {Path(file_path).name}")
        print(f"[Loader] Waiting for Gemini Video processing: {Path(file_path).name}...")
        time.sleep(5)
        video_file = genai.get_file(video_file.name)
        
    if video_file.state.name == "FAILED":
        raise ValueError("Video processing failed in Google Gemini.")
        
    model_name = os.getenv("GOOGLE_MODEL", "gemini-2.5-flash")
    model = genai.GenerativeModel(model_name)
    prompt = "Transcribe the audio accurately and specifically describe the key visual events and slides happening in this video."
    response = model.generate_content([prompt, video_file], request_options={"timeout": 600})
    genai.delete_file(video_file.name)
    return f"--- Video Content ({Path(file_path).name}) ---\n{response.text}"

def extract_text(file_path: str) -> str:
    """
    Extract text from a file based on its extension.
    Supported: .pdf, .docx, .xlsx, .txt, .md, .csv, .png, .jpg, .jpeg, .mp4, .mov
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    extractors = {
        ".pdf": extract_text_from_pdf,
        ".docx": extract_text_from_docx,
        ".xlsx": extract_text_from_xlsx,
        ".xls": extract_text_from_xlsx,
        ".txt": extract_text_from_txt,
        ".md": extract_text_from_txt,
        ".csv": extract_text_from_txt,
        ".png": extract_text_from_image,
        ".jpg": extract_text_from_image,
        ".jpeg": extract_text_from_image,
        ".webp": extract_text_from_image,
        ".heic": extract_text_from_image,
        ".heif": extract_text_from_image,
        ".mp4": extract_text_from_video,
        ".mov": extract_text_from_video,
    }

    extractor = extractors.get(suffix)
    if extractor is None:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: {list(extractors.keys())}")

    return extractor(file_path)


def get_excel_columns(file_path: str) -> List[Dict[str, str]]:
    """
    Get column headers from an Excel file for the column selector UI.
    Returns a list of dicts with 'letter' and 'name' keys.
    """
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    ws = wb.active
    columns = []

    for col_idx, cell in enumerate(next(ws.iter_rows(min_row=1, max_row=1, values_only=False)), start=0):
        col_letter = openpyxl.utils.get_column_letter(col_idx + 1)
        col_name = str(cell.value) if cell.value else f"Column {col_letter}"
        columns.append({"letter": col_letter, "name": col_name})

    wb.close()
    return columns


def get_excel_preview(file_path: str, max_rows: int = 5) -> List[Dict[str, str]]:
    """
    Get a preview of the first N rows from an Excel file.
    Returns a list of dicts mapping column letters to cell values.
    """
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    ws = wb.active
    rows = []

    for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
        if row_idx >= max_rows + 1:  # +1 for header
            break
        row_data = {}
        for col_idx, cell_value in enumerate(row):
            col_letter = openpyxl.utils.get_column_letter(col_idx + 1)
            row_data[col_letter] = str(cell_value) if cell_value is not None else ""
        rows.append(row_data)

    wb.close()
    return rows
